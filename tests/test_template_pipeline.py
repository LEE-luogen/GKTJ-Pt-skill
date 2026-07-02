from __future__ import annotations

import tempfile
import sys
import unittest
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.shared import Pt, RGBColor

FIXTURE_DIR = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(FIXTURE_DIR))

from scripts.build_payload import ANALYSIS_MAX_LEN, ANALYSIS_MIN_LEN, build_payload_v2, parse_markdown_content, text_len, validate_payload_v2
from scripts.render_from_template import render_from_template


TEMPLATE_PATH = FIXTURE_DIR / "templates" / "patient-unified-v1.docx"
MANIFEST_PATH = FIXTURE_DIR / "templates" / "patient-unified-v1.manifest.json"


class Namespace:
    def __init__(self, **kwargs):
        self.__dict__.update(kwargs)


def sample_questionnaire() -> dict:
    question_texts = [
        "您服用该药的主要原因是什么？",
        "您平时监测血压的频率如何？",
        "您对降压作用了解多少？",
        "您是否主动控制饮食与饮酒？",
        "您通常如何获取用药知识？",
        "您是否会自行调整剂量？",
        "您主要通过什么方式提醒自己按时服药？",
        "您是否存在联合用药或保健品使用情况？",
        "您是否愿意参加健康教育活动？",
        "您认为哪些方式最有助于提升依从性？",
    ]
    questions = []
    for index, question_text in enumerate(question_texts, start=1):
        questions.append(
            {
                "number": index,
                "question": question_text,
                "total": "120",
                "options": [
                    {"label": "A", "text": f"选项A{index}", "count": "60", "pct": "50.00%"},
                    {"label": "B", "text": f"选项B{index}", "count": "36", "pct": "30.00%"},
                    {"label": "C", "text": f"选项C{index}", "count": "24", "pct": "20.00%"},
                ],
            }
        )
    return {"questions": questions}


def sample_markdown() -> str:
    chapter2_titles = [
        "用药原因结构分析",
        "血压监测情况分析",
        "降压认知程度分析",
        "饮食控制情况分析",
        "知识获取方式分析",
        "剂量调整风险分析",
        "服药提醒方式分析",
        "联合用药情况分析",
        "健康教育参与分析",
        "依从性提升路径分析",
    ]
    chapter2 = []
    for index, title in enumerate(chapter2_titles, start=1):
        chapter2.append(f"### {index}. {title}\n\n维度{index}分析第一句。维度{index}分析第二句。")
    return """---
product: 测试产品
region: 测试地区
time: 2026.05
survey_period: 2026年5月1日至2026年5月31日
issued_count: 128
valid_count: 120
report_type: 患者端问卷调研分析报告
---

## 一、引言

### 报告背景

第一段引言, 包含"双引号"和(括号)。本次调研从患者视角系统了解用药体验、日常管理行为、信息理解情况和持续支持需求，并结合各维度的选择结构观察不同患者在便利性、规律性、经济负担与沟通反馈方面的差异。报告将问卷信息转化为可用于患者教育、服务优化和健康管理改进的参考，强调如实呈现调研反馈，不将患者感受替代临床证据或专业判断。通过对共性优势、执行难点和支持缺口进行归纳，为后续形成更清晰、更易执行的患者支持方案提供依据。同时关注患者在长期管理中的实际困难，帮助后续服务更准确地回应不同群体的信息、提醒和沟通需求。调研结果只用于归纳患者反馈中的共性趋势，不用于替代个体化的临床评估，并为后续患者服务的持续改进提供线索。

### 数据来源

本报告以测试问卷的有效回收记录为数据基础，按原始题目顺序整理选项、计数和占比，并在生成前检查题目与图表的对应关系。问卷覆盖用药行为、体验评价、信息理解和支持需求等维度，所有比例均依据规范化后的问卷数据计算，不补充未提供的机构、日期或样本信息。数据处理时保留原始问卷题意和选项顺序，对缺失或无法确认的内容仅做记录，不自行推断或补齐。

## 二、数据信息分析

{chapter2}

## 三、反馈意见分析

### 3.1 积极反馈

- 优势一：积极反馈内容一。
- 优势二：积极反馈内容二。
- 优势三：积极反馈内容三。
- 优势四：积极反馈内容四。

### 3.2 待改进反馈

- 问题一：待改进内容一。
- 问题二：待改进内容二。
- 问题三：待改进内容三。
- 问题四：待改进内容四。
- 问题五：待改进内容五。

## 四、综合分析与建议

### 建议一

建议一正文, 包含"示例"提示。

### 建议二

建议二正文。

### 建议三

建议三正文。

### 建议四

建议四正文。

### 建议五

建议五正文。
""".format(chapter2="\n\n".join(chapter2))


class TemplatePipelineTest(unittest.TestCase):
    def build_payload(self) -> dict:
        with tempfile.TemporaryDirectory() as temp_dir:
            report_content = Path(temp_dir) / "content.md"
            report_content.write_text(sample_markdown(), encoding="utf-8")
            meta, content = parse_markdown_content(report_content)
            payload = build_payload_v2(
                sample_questionnaire(),
                meta,
                content,
                Namespace(
                    product=None,
                    region=None,
                    time=None,
                    survey_period=None,
                    sample_size=None,
                    issued_count=None,
                    valid_count=None,
                    report_type=None,
                    data_issue=None,
                    unresolved=None,
                    template_id="patient-unified-v1",
                    template_file="templates/patient-unified-v1.docx",
                ),
            )
        return payload

    def test_build_payload_v2(self) -> None:
        payload = self.build_payload()
        validate_payload_v2(payload)
        self.assertEqual(payload["meta"]["product"], "测试产品")
        self.assertEqual(payload["meta"]["survey_period"], "2026年5月1日至2026年5月31日")
        self.assertEqual(payload["meta"]["issued_count"], "128")
        self.assertEqual(payload["meta"]["valid_count"], "120")
        self.assertEqual(len(payload["chapter2"]["items"]), 10)
        self.assertEqual(payload["chapter2"]["items"][0]["chart_ref"], "chart_slot_01")
        self.assertNotIn("cover", payload)
        self.assertTrue(4 <= len(payload["chapter2"]["items"][0]["title"]) <= 8)
        self.assertTrue(ANALYSIS_MIN_LEN <= text_len(payload["chapter2"]["items"][0]["analysis"]) <= ANALYSIS_MAX_LEN)
        self.assertEqual(payload["summary"]["recommendations"][0]["title"], "建议一")

    def test_render_from_template_uses_body_only_layout_and_image_charts(self) -> None:
        payload = self.build_payload()
        with tempfile.TemporaryDirectory() as temp_dir:
            output_docx = Path(temp_dir) / "rendered.docx"
            render_from_template(payload, TEMPLATE_PATH, MANIFEST_PATH, output_docx)

            document = Document(output_docx)
            texts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            all_text = "\n".join(texts)

            self.assertEqual(texts[0], "目录（请在 Word 中按 Ctrl+A, F9 更新）")
            self.assertIn("引言", texts)
            self.assertIn("报告背景", texts)
            self.assertIn("数据来源", texts)
            self.assertIn("数据信息分析", texts)
            self.assertIn("综合分析与建议", texts)
            self.assertIn("附件-问卷题目内容", texts)
            self.assertNotIn("问卷调研分析报告", texts)
            self.assertNotIn("QUESTIONNAIRE SURVEY ANALYSIS", all_text)
            self.assertNotIn("国开（天津）信息科技有限公司", all_text)
            self.assertNotIn("占比：", all_text)
            self.assertIn('第一段引言, 包含"双引号"和(括号)。', all_text)
            self.assertNotIn("### 报告背景", all_text)
            self.assertIn('建议一正文, 包含"示例"提示。', all_text)
            self.assertEqual(len(document.inline_shapes), 11)

            question_title = next(paragraph for paragraph in document.paragraphs if paragraph.text.strip() == "用药原因结构分析")
            self.assertEqual(question_title.runs[0].font.name, "宋体")
            self.assertEqual(question_title.runs[0].font.size, Pt(16))
            self.assertTrue(question_title.runs[0].bold)

            summary_title = next(paragraph for paragraph in document.paragraphs if paragraph.text.strip() == "建议一")
            self.assertEqual(summary_title.runs[0].font.name, "宋体")
            self.assertEqual(summary_title.runs[0].font.size, Pt(16))
            self.assertTrue(summary_title.runs[0].bold)

            body_paragraph = next(paragraph for paragraph in document.paragraphs if paragraph.text.strip().startswith("第一段引言"))
            template_body = Document(TEMPLATE_PATH).paragraphs[47]
            self.assertEqual(body_paragraph.paragraph_format.first_line_indent, template_body.paragraph_format.first_line_indent)
            self.assertEqual(body_paragraph.paragraph_format.space_before, template_body.paragraph_format.space_before)
            self.assertEqual(body_paragraph.paragraph_format.space_after, template_body.paragraph_format.space_after)
            self.assertEqual(body_paragraph.paragraph_format.line_spacing, template_body.paragraph_format.line_spacing)
            self.assertEqual(body_paragraph.alignment, template_body.alignment)

            attachment_question = next(paragraph for paragraph in document.paragraphs if paragraph.text.strip() == "1.您服用该药的主要原因是什么？")
            self.assertEqual(attachment_question.runs[0].font.name, "宋体")
            self.assertEqual(attachment_question.runs[0].font.size, Pt(16))


if __name__ == "__main__":
    unittest.main()
