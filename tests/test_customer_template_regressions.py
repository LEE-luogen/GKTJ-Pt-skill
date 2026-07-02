import copy
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt
from lxml import etree

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from scripts.build_payload import question_options
from scripts.render_from_template import render_from_template


TEMPLATE = ROOT / "templates" / "patient-unified-v1.docx"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
C_NS = "http://schemas.openxmlformats.org/drawingml/2006/chart"


def make_payload() -> dict:
    items = []
    questions = []
    for index in range(1, 12):
        options = [
            {"code": "A", "text": f"选项A{index}", "pct": 60.25},
            {"code": "B", "text": f"选项B{index}", "pct": 39.75},
        ]
        items.append(
            {
                "question_ref": f"q{index:02d}",
                "chart_ref": f"chart_slot_{index:02d}",
                "title": f"维度{index:02d}",
                "analysis": f"第{index}道题分析正文，60.25%的患者选择选项A。",
                "chart": {"options": copy.deepcopy(options)},
            }
        )
        questions.append(
            {
                "number": index,
                "question_ref": f"q{index:02d}",
                "question": f"第{index}道问卷题目？",
                "options": copy.deepcopy(options),
            }
        )
    return {
        "meta": {"product": "测试产品", "region": "河南省"},
        "intro": {"background": "背景" * 160, "data_source": "数据来源" * 40},
        "chapter2": {"items": items},
        "feedback": {
            "positive": [{"title": "积极标题", "body": "积极反馈正文。"}],
            "negative": [{"title": "改进标题", "body": "改进反馈正文。"}],
        },
        "summary": {
            "recommendations": [
                {"title": f"建议标题{index}", "body": f"建议正文{index}。"}
                for index in range(1, 5)
            ]
        },
        "attachment": {"questions": questions},
    }


def test_question_options_accepts_code_percentage_schema():
    question = {
        "options": [
            {"code": "A", "text": "A.短时间内起效", "count": 31, "percentage": 30.98}
        ]
    }
    assert question_options(question) == [
        {"label": "A", "text": "短时间内起效", "pct": "30.98%"}
    ]


def test_render_preserves_charts_and_chapter_boundaries(tmp_path):
    output = tmp_path / "report.docx"
    render_from_template(make_payload(), TEMPLATE, output)

    doc = Document(output)
    texts = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    assert len(doc.inline_shapes) == 11
    assert texts.index("目录（请在 Word 中按 Ctrl+A, F9 更新）") < texts.index("引言")
    assert "附件-问卷题目内容" in texts
    attachment_index = texts.index("附件-问卷题目内容")
    assert texts[attachment_index + 1] == "1.第1道问卷题目？"

    summary_index = texts.index("综合分析与建议")
    assert texts[summary_index + 1 : summary_index + 9] == [
        "建议标题1", "建议正文1。",
        "建议标题2", "建议正文2。",
        "建议标题3", "建议正文3。",
        "建议标题4", "建议正文4。",
    ]

    title_para = next(p for p in doc.paragraphs if p.text.strip() == "建议标题1")
    body_para = next(p for p in doc.paragraphs if p.text.strip() == "建议正文1。")
    assert title_para.runs[0].font.size == Pt(16)
    assert title_para.runs[0].bold is True
    assert body_para.runs[0].font.size == Pt(14)
    assert body_para.runs[0].bold is not True

    dimension_heading = next(p for p in doc.paragraphs if p.text.strip() == "维度01")
    assert dimension_heading._p.pPr.numPr is not None
    assert not dimension_heading.text.startswith("·")

    for heading in ("引言", "数据信息分析", "反馈意见分析", "综合分析与建议", "附件-问卷题目内容"):
        paragraph = next(p for p in doc.paragraphs if p.text.strip() == heading)
        assert paragraph._p.pPr.numPr is not None
        assert not paragraph.text[:2] in {"一、", "二、", "三、", "四、", "五、"}

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            r_fonts = run._r.get_or_add_rPr().get_or_add_rFonts()
            assert r_fonts.get(f"{{{W_NS}}}ascii") == "宋体"
            assert r_fonts.get(f"{{{W_NS}}}hAnsi") == "宋体"
            assert r_fonts.get(f"{{{W_NS}}}eastAsia") == "宋体"

    with zipfile.ZipFile(output) as zf:
        chart_names = sorted(
            name for name in zf.namelist()
            if name.startswith("word/charts/chart") and name.endswith(".xml")
        )
        assert len(chart_names) == 11
        first_chart = etree.fromstring(zf.read(chart_names[0]))
        values = first_chart.xpath("//c:ser/c:val//c:pt/c:v/text()", namespaces={"c": C_NS})
        assert values == ["0.6025", "0.3975"]
        chart_fonts = first_chart.xpath("//a:defRPr", namespaces={"a": "http://schemas.openxmlformats.org/drawingml/2006/main"})
        assert chart_fonts
        assert {node.get("sz") for node in chart_fonts} == {"1000"}
        assert all(node.find(f"{{http://schemas.openxmlformats.org/drawingml/2006/main}}ea").get("typeface") == "宋体" for node in chart_fonts)
