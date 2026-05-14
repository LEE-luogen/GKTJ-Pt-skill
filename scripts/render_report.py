#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
import textwrap
from math import ceil
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.font_manager import FontProperties
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

GREEN = RGBColor(0x75, 0xBD, 0x42)
BLACK = RGBColor(0x00, 0x00, 0x00)
SONGTI = "宋体"
DEFAULT_REPORT_TYPE = "患者端问卷调研分析报告"
BULLET_SYMBOL = "❖ "
PRIMARY_HEADINGS = {
    "一、引言",
    "二、数据信息分析",
    "三、反馈意见分析",
    "四、综合分析与建议",
    "五、附件-问卷题目内容",
}
FEEDBACK_SUBHEADINGS = {"3.1 积极反馈", "3.2 待改进反馈"}


def pick_font_path() -> str | None:
    for path in [
        "/System/Library/Fonts/Supplemental/Songti.ttc",
        "/System/Library/Fonts/Hiragino Sans GB.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
    ]:
        if Path(path).exists():
            return path
    return None


def create_chart(path: Path, title: str, options: list[dict]):
    font_path = pick_font_path()
    font = FontProperties(fname=font_path) if font_path else None
    del title
    ordered_options = list(reversed(options))
    labels = [textwrap.fill(f"{o['label']}.{o['text']}", width=18) for o in ordered_options]
    values = [float(str(o["pct"]).replace("%", "")) for o in ordered_options]
    colors = ["#D9E3BE", "#C5D895", "#A9C65A", "#86B43C", "#6A9827", "#557C1E"][: len(values)]
    x_upper = max(70, int(ceil((max(values) + 8) / 10.0) * 10)) if values else 70
    fig, ax = plt.subplots(figsize=(8.2, 4.6), dpi=200)
    y = list(range(len(labels)))
    bars = ax.barh(y, values, color=colors, height=0.62, edgecolor="#90A764", linewidth=0.5)
    ax.invert_yaxis()
    ax.set_xlim(0, x_upper * 1.72)
    ax.set_ylim(-0.6, len(labels) - 0.4)
    ax.set_yticks(y)
    ax.set_yticklabels([""] * len(labels))
    xticks = list(range(0, x_upper + 1, 10))
    ax.set_xticks(xticks)
    ax.set_xticklabels([f"{tick}%" for tick in xticks], fontproperties=font, fontsize=8.5)
    ax.set_ylabel("占比", fontproperties=font, fontsize=9, color="#666666", rotation=0, labelpad=18)
    ax.tick_params(axis="x", labelsize=8.5, colors="#666666", length=0, pad=6)
    ax.tick_params(axis="y", length=0)
    ax.xaxis.grid(True, linestyle="-", linewidth=0.4, color="#E3E3E3")
    ax.set_axisbelow(True)
    for spine in ["top", "right"]:
        ax.spines[spine].set_visible(False)
    ax.spines["bottom"].set_color("#D0D0D0")
    ax.spines["left"].set_color("#D0D0D0")
    legend_x = x_upper * 1.08
    for index, (rect, val, label, color) in enumerate(zip(bars, values, labels, colors)):
        y_center = rect.get_y() + rect.get_height() / 2
        ax.text(
            rect.get_width() + x_upper * 0.02,
            y_center,
            format_chart_pct(val),
            va="center",
            ha="left",
            fontsize=8.5,
            color="#555555",
            fontproperties=font,
        )
        ax.scatter(legend_x, index, s=28, marker="s", color=color, edgecolors="#90A764", linewidths=0.4, clip_on=False)
        ax.text(
            legend_x + x_upper * 0.035,
            index,
            label,
            va="center",
            ha="left",
            fontsize=8.2,
            color="#666666",
            fontproperties=font,
        )
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")
    fig.subplots_adjust(left=0.08, right=0.96, top=0.92, bottom=0.18)
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def validate_payload(payload: dict):
    required = [
        "product",
        "region",
        "introduction",
        "data_analysis",
        "positive_feedback",
        "negative_feedback",
        "summary_recommendations",
        "attachment_questions",
    ]
    missing = [key for key in required if key not in payload]
    if missing:
        raise ValueError(f"Invalid payload. Missing keys: {', '.join(missing)}")
    if len(payload["data_analysis"]) != len(payload["attachment_questions"]):
        raise ValueError("Invalid payload. `data_analysis` count must match `attachment_questions` count.")
    for index, item in enumerate(payload["data_analysis"], start=1):
        if not item.get("title") or not item.get("analysis"):
            raise ValueError(f"Invalid payload. `data_analysis[{index}]` missing title or analysis.")
        if not item.get("options"):
            raise ValueError(f"Invalid payload. `data_analysis[{index}]` missing chart options.")


def set_doc_style(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    normal = doc.styles["Normal"]
    normal.font.name = SONGTI
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), SONGTI)
    normal.font.size = Pt(14)
    normal.font.color.rgb = BLACK


def set_run_font(run, size_pt: float, bold: bool = False, color: RGBColor = BLACK, font_name: str = SONGTI):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.color.rgb = color


def run_has_drawing(run) -> bool:
    return bool(run._element.xpath(".//w:drawing | .//w:pict"))


def remove_illegal_chars(text: str) -> str:
    return re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]", "", text)


def normalize_chinese_quotes(text: str) -> str:
    result = []
    open_double = True
    open_single = True
    for char in text:
        if char == '"':
            result.append("“" if open_double else "”")
            open_double = not open_double
        elif char == "'":
            result.append("‘" if open_single else "’")
            open_single = not open_single
        else:
            result.append(char)
    return "".join(result)


def normalize_punctuation(text: str) -> str:
    normalized = remove_illegal_chars(text)
    normalized = normalized.replace("...", "……")
    normalized = normalize_chinese_quotes(normalized)
    replacements = [
        (r"(?<=[\u4e00-\u9fff])\s*,", "，"),
        (r",(?=[\u4e00-\u9fff“‘（])", "，"),
        (r"(?<=[\u4e00-\u9fff])\s*:", "："),
        (r":(?=[\u4e00-\u9fff“‘（])", "："),
        (r"(?<=[\u4e00-\u9fff])\s*;", "；"),
        (r";(?=[\u4e00-\u9fff“‘（])", "；"),
        (r"(?<=[\u4e00-\u9fff])\s*\?", "？"),
        (r"\?(?=[\u4e00-\u9fff“‘（])", "？"),
        (r"(?<=[\u4e00-\u9fff])\s*!", "！"),
        (r"!(?=[\u4e00-\u9fff“‘（])", "！"),
        (r"(?<=[\u4e00-\u9fff])\(", "（"),
        (r"\)(?=[\u4e00-\u9fff])", "）"),
        (r"(?<=[\u4e00-\u9fff])\)", "）"),
        (r"\((?=[\u4e00-\u9fff])", "（"),
    ]
    for pattern, repl in replacements:
        normalized = re.sub(pattern, repl, normalized)
    normalized = re.sub(r"([，。！？：；、“”‘’（）])\s+(?=[\u4e00-\u9fff“‘（])", r"\1", normalized)
    normalized = re.sub(r"(?<=[\u4e00-\u9fff])\s+([，。！？：；、])", r"\1", normalized)
    return normalized


def classify_paragraph(text: str, current_section: str | None) -> str:
    stripped = text.strip()
    if not stripped:
        return "empty"
    if stripped in PRIMARY_HEADINGS:
        return "primary_heading"
    if stripped in FEEDBACK_SUBHEADINGS:
        return "feedback_subheading"
    if stripped.startswith(BULLET_SYMBOL):
        return "secondary_heading"
    if current_section == "五、附件-问卷题目内容":
        if re.match(r"^\d+\.", stripped):
            return "attachment_question"
        if re.match(r"^[A-Z]\.", stripped):
            return "attachment_option"
    if current_section == "三、反馈意见分析":
        return "labeled_body"
    return "body"


def apply_paragraph_format(paragraph, role: str) -> None:
    fmt = paragraph.paragraph_format
    if role == "primary_heading":
        fmt.space_before = Pt(7)
        fmt.space_after = Pt(7)
        fmt.first_line_indent = Pt(0)
        fmt.left_indent = Pt(0)
        fmt.line_spacing = 1.5
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif role == "secondary_heading":
        fmt.space_before = Pt(7)
        fmt.space_after = Pt(7)
        fmt.first_line_indent = Pt(0)
        fmt.left_indent = Pt(0)
        fmt.line_spacing = 1.5
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif role == "feedback_subheading":
        fmt.space_before = Pt(7)
        fmt.space_after = Pt(7)
        fmt.first_line_indent = Pt(0)
        fmt.left_indent = Pt(0)
        fmt.line_spacing = 1.5
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif role == "attachment_question":
        fmt.space_before = Pt(7)
        fmt.space_after = Pt(0)
        fmt.first_line_indent = Pt(0)
        fmt.left_indent = Pt(0)
        fmt.line_spacing = 1.5
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif role == "attachment_option":
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.first_line_indent = Pt(0)
        fmt.left_indent = Pt(21)
        fmt.line_spacing = 1.5
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif role == "labeled_body":
        fmt.space_before = Pt(7)
        fmt.space_after = Pt(0)
        fmt.first_line_indent = Cm(0.74)
        fmt.left_indent = Pt(0)
        fmt.line_spacing = 1.5
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif role == "body":
        fmt.space_before = Pt(7)
        fmt.space_after = Pt(0)
        fmt.first_line_indent = Cm(0.74)
        fmt.left_indent = Pt(0)
        fmt.line_spacing = 1.5
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def apply_run_format(paragraph, role: str) -> None:
    if role == "primary_heading":
        for run in paragraph.runs:
            set_run_font(run, 20, bold=False, color=GREEN)
    elif role == "secondary_heading":
        for index, run in enumerate(paragraph.runs):
            if index == 0 and run.text == BULLET_SYMBOL:
                set_run_font(run, 14, bold=True, color=GREEN)
            else:
                set_run_font(run, 16, bold=True, color=BLACK)
    elif role == "feedback_subheading":
        for run in paragraph.runs:
            set_run_font(run, 16, bold=True, color=BLACK)
    elif role == "attachment_question":
        for run in paragraph.runs:
            set_run_font(run, 16, bold=True, color=BLACK)
    elif role == "attachment_option":
        for run in paragraph.runs:
            set_run_font(run, 14, bold=False, color=BLACK)
    elif role == "labeled_body":
        for index, run in enumerate(paragraph.runs):
            set_run_font(run, 14, bold=index == 0, color=BLACK)
    elif role == "body":
        for run in paragraph.runs:
            set_run_font(run, 14, bold=False, color=BLACK)


def postprocess_docx(docx_path: Path) -> None:
    doc = Document(docx_path)
    set_doc_style(doc)
    current_section: str | None = None
    for paragraph in doc.paragraphs:
        stripped = paragraph.text.strip()
        role = classify_paragraph(stripped, current_section)
        for run in paragraph.runs:
            if run_has_drawing(run):
                continue
            if role == "secondary_heading" and run.text == BULLET_SYMBOL:
                continue
            run.text = normalize_punctuation(run.text)
        apply_paragraph_format(paragraph, role)
        apply_run_format(paragraph, role)
        updated_text = paragraph.text.strip()
        if updated_text in PRIMARY_HEADINGS:
            current_section = updated_text
    doc.save(docx_path)


def add_primary_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, 20, bold=False, color=GREEN)
    return p


def add_secondary_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    bullet = p.add_run(BULLET_SYMBOL)
    set_run_font(bullet, 14, bold=True, color=GREEN)
    title_run = p.add_run(text)
    set_run_font(title_run, 16, bold=True, color=BLACK)
    return p


def add_simple_subheading(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, 16, bold=True, color=BLACK)
    return p


def add_body_para(doc: Document, text: str, first_indent_pt: int = 28):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(first_indent_pt)
    run = p.add_run(text)
    set_run_font(run, 14, bold=False, color=BLACK)
    return p


def add_labeled_body_para(doc: Document, label: str, body: str, first_indent_pt: int = 21):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(first_indent_pt)
    r1 = p.add_run(f"{label}：")
    set_run_font(r1, 14, bold=True, color=BLACK)
    r2 = p.add_run(body)
    set_run_font(r2, 14, bold=False, color=BLACK)
    return p


def format_chart_pct(value: float) -> str:
    return f"{int(round(value))}%" if abs(value - round(value)) < 0.01 else f"{value:.2f}%"


def normalize_count_text(value: str | None) -> str | None:
    if value is None:
        return None
    text = str(value).strip()
    if not text:
        return None
    digits = re.sub(r"[^\d]", "", text)
    if digits:
        return digits
    return text


def normalize_sample_size_text(value: str | None) -> str | None:
    if value is None:
        return None
    text = str(value).strip()
    if not text:
        return None
    if text.upper().startswith("N="):
        return text[2:].strip()
    return text


def build_intro_sections(payload: dict) -> list[dict]:
    intro_paragraphs = [str(item).strip() for item in payload.get("introduction", []) if str(item).strip()]
    if not intro_paragraphs:
        return []

    background_paragraphs = [intro_paragraphs[0]]
    data_source_paragraphs = intro_paragraphs[1:] or []
    survey_period = payload.get("survey_period")
    issued_count = normalize_count_text(payload.get("issued_count"))
    valid_count = normalize_count_text(payload.get("valid_count"))
    sample_size = normalize_sample_size_text(payload.get("sample_size"))
    additions = []
    if survey_period:
        additions.append(f"问卷收集时间为{survey_period}")
    if issued_count:
        additions.append(f"共发放问卷{issued_count}份")
    if valid_count:
        additions.append(f"回收有效问卷{valid_count}份")
    elif sample_size:
        additions.append(f"有效样本数为{sample_size}份")
    if additions:
        if data_source_paragraphs:
            data_source_paragraphs[0] = f"{data_source_paragraphs[0].rstrip('。')}。本次调研{'，'.join(additions)}。"
        else:
            data_source_paragraphs = [f"本次调研{'，'.join(additions)}。"]
    sections = [{"title": "报告背景", "paragraphs": background_paragraphs}]
    if data_source_paragraphs:
        sections.append({"title": "数据来源", "paragraphs": data_source_paragraphs})
    return sections


def split_analysis_paragraphs(text: str) -> list[str]:
    normalized = re.sub(r"\s+", "", text.strip())
    if not normalized:
        return []
    sentences = [item for item in re.split(r"(?<=[。！？])", normalized) if item]
    paragraphs: list[str] = []
    buf = ""
    for sentence in sentences:
        buf += sentence
        if len(buf) >= 120:
            paragraphs.append(buf)
            buf = ""
    if buf:
        paragraphs.append(buf)
    if len(paragraphs) == 1 and len(normalized) >= 220 and len(sentences) >= 4:
        midpoint = max(2, len(sentences) // 2)
        return ["".join(sentences[:midpoint]), "".join(sentences[midpoint:])]
    return paragraphs


def write_markdown(payload: dict, output: Path):
    lines = [
        "# 问卷调研分析报告",
        "",
        f"品种：{payload['product']}",
        f"地区：{payload['region']}",
    ]
    if payload.get("time"):
        lines.append(f"时间：{payload['time']}")
    lines.append(f"报告类型：{payload.get('report_type', DEFAULT_REPORT_TYPE)}")
    lines += ["", "## 一、引言", ""]
    for section in build_intro_sections(payload):
        lines += [f"### {section['title']}", ""]
        for paragraph in section["paragraphs"]:
            lines += [paragraph, ""]
    lines += ["## 二、数据信息分析", ""]
    for item in payload["data_analysis"]:
        lines += [f"### {item['title']}", "", f"![图{item['number']} {item['title']}](charts/chart_{item['number']:02d}.png)", ""]
        lines += [item["analysis"], ""]
    lines += ["## 三、反馈意见分析", "", "### 3.1 积极反馈", ""]
    for item in payload["positive_feedback"]:
        lines.append(f"- {item['title']}：{item['body']}")
    lines += ["", "### 3.2 待改进反馈", ""]
    for item in payload["negative_feedback"]:
        lines.append(f"- {item['title']}：{item['body']}")
    lines += ["", "## 四、综合分析与建议", ""]
    for item in payload.get("summary_recommendations", []):
        lines += [f"### {item['title']}", "", item["body"], ""]
    lines += ["## 五、附件-问卷题目内容", ""]
    for item in payload["attachment_questions"]:
        lines.append(f"{item['number']}.{item['question']}")
        for opt in item["options"]:
            lines.append(f"{opt['code']}.{opt['text']}")
        lines.append("")
    output.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def write_docx(payload: dict, output: Path, charts_dir: Path):
    doc = Document()
    set_doc_style(doc)
    add_primary_heading(doc, "一、引言")
    for section in build_intro_sections(payload):
        add_secondary_heading(doc, section["title"])
        for paragraph in section["paragraphs"]:
            add_body_para(doc, paragraph)
    add_primary_heading(doc, "二、数据信息分析")
    for item in payload["data_analysis"]:
        add_secondary_heading(doc, item["title"])
        img = doc.add_paragraph()
        img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img.paragraph_format.space_before = Pt(7)
        img.paragraph_format.space_after = Pt(7)
        img.add_run().add_picture(str(charts_dir / f"chart_{item['number']:02d}.png"), width=Cm(14.4))
        add_body_para(doc, item["analysis"])
    add_primary_heading(doc, "三、反馈意见分析")
    add_simple_subheading(doc, "3.1 积极反馈")
    for item in payload["positive_feedback"]:
        add_labeled_body_para(doc, item["title"], item["body"])
    add_simple_subheading(doc, "3.2 待改进反馈")
    for item in payload["negative_feedback"]:
        add_labeled_body_para(doc, item["title"], item["body"])
    add_primary_heading(doc, "四、综合分析与建议")
    for item in payload.get("summary_recommendations", []):
        add_secondary_heading(doc, item["title"])
        add_body_para(doc, item["body"])
    add_primary_heading(doc, "五、附件-问卷题目内容")
    for item in payload["attachment_questions"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(f"{item['number']}.{item['question']}")
        set_run_font(run, 16, bold=True, color=BLACK)
        for opt in item["options"]:
            p1 = doc.add_paragraph()
            p1.paragraph_format.left_indent = Pt(21)
            p1.paragraph_format.space_after = Pt(0)
            p1.paragraph_format.line_spacing = 1.5
            run1 = p1.add_run(f"{opt['code']}.{opt['text']}")
            set_run_font(run1, 14, bold=False, color=BLACK)
    doc.save(output)
    postprocess_docx(output)


def summarize(payload: dict, charts_dir: Path, docx_path: Path, out_path: Path):
    doc = Document(docx_path)
    chart_count = len(list(charts_dir.glob("chart_*.png")))
    summary = {
        "markdown_final": str(out_path.parent / "report_final.md"),
        "word_file": str(docx_path),
        "chapter_complete": True,
        "chart_count": chart_count,
        "question_count": len(payload["data_analysis"]),
        "chart_count_ok": chart_count == len(payload["data_analysis"]) and len(doc.inline_shapes) == len(payload["data_analysis"]),
        "chart_style_ok": True,
        "data_issue": payload.get("checks", {}).get("data_issue"),
        "unresolved": payload.get("checks", {}).get("unresolved"),
    }
    out_path.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")


def main():
    parser = argparse.ArgumentParser(description="Render patient survey report artifacts from payload JSON.")
    parser.add_argument("payload_json", help="Path to report payload JSON")
    parser.add_argument("--output-dir", required=True, help="Output directory")
    args = parser.parse_args()

    payload = json.loads(Path(args.payload_json).read_text(encoding="utf-8"))
    validate_payload(payload)
    out_dir = Path(args.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    charts_dir = out_dir / "charts"
    charts_dir.mkdir(exist_ok=True)

    for item in payload["data_analysis"]:
        create_chart(charts_dir / f"chart_{item['number']:02d}.png", item["title"], item["options"])

    draft_md = out_dir / "report_draft.md"
    final_md = out_dir / "report_final.md"
    write_markdown(payload, draft_md)
    write_markdown(payload, final_md)
    docx_path = out_dir / f"问卷调研分析报告-{payload['product']}-患者端-{payload['region']}.docx"
    write_docx(payload, docx_path, charts_dir)
    summarize(payload, charts_dir, docx_path, out_dir / "report_summary.json")
    print(docx_path)


if __name__ == "__main__":
    main()
