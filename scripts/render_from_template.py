#!/usr/bin/env python3
"""Render a patient questionnaire report using the customer DOCX template.

This is the main entry point for the v2 rendering pipeline. It:
1. Copies the customer DOCX template.
2. Replaces cover, intro, chapters 2-5 text content.
3. Preserves native Word charts (updated separately by update_word_charts.py).
4. Inserts a TOC field after the cover section.
5. Sets updateFields=true in settings.xml.
6. Validates the output.

Usage:
    python render_from_template.py <payload_v2.json> -o <output.docx>
    python render_from_template.py <payload_v2.json> -o <output.docx> --template <template.docx>
"""
from __future__ import annotations

import argparse
import copy
import json
import shutil
import sys
import zipfile
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor
from lxml import etree

CURRENT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = CURRENT_DIR.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from scripts.update_word_charts import (
    ChartSlotData,
    payload_to_chart_slots,
    update_charts_in_docx,
)

# ── Constants ────────────────────────────────────────────────────────────────
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
GREEN_HEX = "2E7D32"  # Template green color for headings
BLACK = RGBColor(0, 0, 0)

# Paragraph indices in the template (from manifest analysis)
COVER_PARA_END = 44       # Customer template cover ends before paragraph 45 (引言)
TOC_INSERT_AFTER = 44     # Insert TOC between cover and introduction

# Chapter paragraph ranges (inclusive)
PRIMARY_HEADINGS = {
    # The visible Chinese chapter numbers come from the template's w:numPr.
    # Keep them out of text or Word renders "一、一、引言".
    "引言": "引言",
    "数据信息分析": "数据信息分析",
    "反馈意见分析": "反馈意见分析",
    "综合分析与建议": "综合分析与建议",
    "附件-问卷题目内容": "附件-问卷题目内容",
}


# ── Text helpers ─────────────────────────────────────────────────────────────

def _text_len(text: str) -> int:
    """Count Chinese characters excluding whitespace."""
    return len(text.replace(" ", "").replace("\n", "").replace("\t", "").replace("\r", ""))


def _replace_paragraph_text(para, new_text: str, font_size_pt: int = 14, bold: bool = False, color=None):
    """Replace all text in a paragraph, preserving paragraph properties.

    For body paragraphs: clear runs, add one new run with the text.
    For headings: same, but with heading formatting.
    For chart paragraphs: skip text replacement to preserve the chart.
    """
    # Skip paragraphs with drawings (charts)
    if _has_drawing(para):
        return

    # Clear existing runs
    for run in para.runs:
        run._element.getparent().remove(run._element)

    # Add new run
    run = para.add_run(new_text)
    run.font.size = Pt(font_size_pt)
    run.font.name = "宋体"
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _replace_paragraph_text_preserve_format(para, new_text: str):
    """Replace text content while preserving existing run formatting.

    For paragraphs with drawings (charts), only replace text in non-drawing runs.
    For plain text paragraphs, clear runs and add one new run.
    """
    # Check if paragraph has a drawing element
    has_chart = _has_drawing(para)

    if has_chart:
        text_nodes = []
        for run in para.runs:
            if not run._element.findall(f".//{{{W_NS}}}drawing"):
                text_nodes.extend(run._element.findall(f".//{{{W_NS}}}t"))
        if text_nodes:
            text_nodes[0].text = new_text
            for node in text_nodes[1:]:
                node.text = ""
        return

    # For non-chart paragraphs, replace all text
    if not para.runs:
        run = para.add_run(new_text)
        run.font.size = Pt(14)
        run.font.name = "宋体"
        return

    # Get formatting from first run
    first_run = para.runs[0]
    font_size = first_run.font.size
    font_name = first_run.font.name
    bold = first_run.bold

    # Clear all runs
    for run in para.runs:
        run._element.getparent().remove(run._element)

    # Add new run preserving format
    run = para.add_run(new_text)
    if font_size:
        run.font.size = font_size
    if font_name:
        run.font.name = font_name
    if bold:
        run.bold = bold


def _has_drawing(para) -> bool:
    """Check if a paragraph contains a drawing element (chart/image)."""
    return bool(para._element.findall(f".//{{{W_NS}}}drawing"))


# ── Cover replacement ────────────────────────────────────────────────────────

def _replace_cover(doc: Document, meta: dict):
    """Replace cover page text (product, disease, region, time, company)."""
    product = meta.get("product", "")
    disease = meta.get("disease", "")
    region = meta.get("region", "")
    report_year = meta.get("report_year", "")
    report_month = meta.get("report_month", "")
    company = meta.get("company", "")

    # Build the report subtitle
    subtitle = f"{product}对特定疾病领域的调研问卷调研分析报告"
    if disease:
        subtitle = f"{product}对{disease}的调研问卷调研分析报告"

    # Replace text in cover paragraphs
    # Para[0]: main cover drawing (skip - it's a complex drawing element)
    # Para[1]: report title
    # Para[6]: subtitle
    # Para[13]: company + date
    for i, para in enumerate(doc.paragraphs):
        if i > COVER_PARA_END:
            break
        text = para.text.strip()
        if not text:
            continue

        # Report title paragraphs
        if "问卷调研分析报告" in text and len(text) < 20:
            _replace_paragraph_text_preserve_format(para, "问卷调研分析报告")

        # Subtitle with product/disease info
        if "厄贝沙坦氢氯噻嗪片" in text and "调研" in text:
            new_text = text.replace("厄贝沙坦氢氯噻嗪片", product)
            if disease and "特定疾病领域" in new_text:
                new_text = new_text.replace("特定疾病领域", disease)
            _replace_paragraph_text_preserve_format(para, new_text)

        # Region
        if "新疆维吾尔自治区" in text:
            new_text = text.replace("新疆维吾尔自治区", region)
            _replace_paragraph_text_preserve_format(para, new_text)

        # Company + date
        if "信息科技有限公司" in text:
            date_str = f"{report_year}.{report_month}" if report_year and report_month else ""
            new_text = f"{company}{date_str}"
            _replace_paragraph_text_preserve_format(para, new_text)


# ── Chapter 1: Introduction ──────────────────────────────────────────────────

def _replace_chapter1(doc: Document, intro: dict):
    """Replace 引言 section: 报告背景 + 数据来源."""
    paragraphs = doc.paragraphs
    background = intro.get("background", "")
    data_source = intro.get("data_source", "")

    # Customer template: 45 引言, 46 报告背景, 47 正文, 48 数据来源, 49 正文.
    if len(paragraphs) > 47 and background:
        _replace_paragraph_text_preserve_format(paragraphs[47], background)

    # Para[63]: data source text
    if len(paragraphs) > 49 and data_source:
        _replace_paragraph_text_preserve_format(paragraphs[49], data_source)


# ── Chapter 2: Data Analysis ─────────────────────────────────────────────────

def _replace_chapter2(doc: Document, chapter2: dict):
    """Replace 数据信息分析 section: headings + analysis text. Charts are untouched."""
    items = chapter2.get("items", [])
    paragraphs = doc.paragraphs

    # Chart slot mapping from manifest:
    # heading_para, chart_para, analysis_para
    slot_map = [
        (51, 52, 53), (54, 55, 56), (57, 58, 59), (60, 61, 62),
        (63, 64, 65), (66, 67, 68), (69, 69, 70), (71, 71, 72),
        (73, 74, 75), (76, 76, 77), (78, 78, 79),
    ]

    for idx, item in enumerate(items):
        if idx >= len(slot_map):
            break

        heading_para_idx, chart_para_idx, analysis_para_idx = slot_map[idx]
        title = item.get("title", "")
        analysis = item.get("analysis", "")

        # Replace heading text (only if it's not a chart paragraph)
        if heading_para_idx < len(paragraphs):
            para = paragraphs[heading_para_idx]
            if _has_drawing(para):
                _replace_paragraph_text_preserve_format(para, title)
            else:
                # The template paragraph already owns the diamond bullet via
                # w:numPr. Adding a literal bullet duplicates it in Word's
                # navigation pane and visible heading.
                _replace_paragraph_text(para, title, font_size_pt=16, bold=True)

        # Replace analysis text
        if analysis_para_idx < len(paragraphs):
            para = paragraphs[analysis_para_idx]
            if not _has_drawing(para):
                _replace_paragraph_text_preserve_format(para, analysis)

        # For shared heading+chart paragraphs, we can't replace text without
        # losing the chart. The title is already in the template. If the title
        # needs to change, we'd need to manipulate the XML directly.
        # For now, we skip text replacement on chart-bearing paragraphs.


# ── Chapter 3: Feedback ──────────────────────────────────────────────────────

def _replace_chapter3(doc: Document, feedback: dict):
    """Replace 反馈意见分析 section."""
    paragraphs = doc.paragraphs
    positive = feedback.get("positive", [])
    negative = feedback.get("negative", [])

    def replace_feedback(paragraph, item):
        title = str(item.get("title", "")).strip()
        body = str(item.get("body", "")).strip()
        for run in paragraph.runs:
            run._element.getparent().remove(run._element)
        title_run = paragraph.add_run(f"{title}：" if title else "")
        title_run.font.size = Pt(14)
        title_run.font.name = "宋体"
        title_run.bold = True
        body_run = paragraph.add_run(body)
        body_run.font.size = Pt(14)
        body_run.font.name = "宋体"
        body_run.bold = False

    # Customer template: positive 83-86, negative 88-91.
    for idx, item in enumerate(positive):
        para_idx = 83 + idx
        if para_idx <= 86 and para_idx < len(paragraphs):
            replace_feedback(paragraphs[para_idx], item)

    # Negative feedback: paras 102-105 (4 items in template)
    for idx, item in enumerate(negative):
        para_idx = 88 + idx
        if para_idx <= 91 and para_idx < len(paragraphs):
            replace_feedback(paragraphs[para_idx], item)


# ── Chapter 4: Summary ───────────────────────────────────────────────────────

def _replace_chapter4(doc: Document, summary: dict):
    """Replace 综合分析与建议 section."""
    paragraphs = doc.paragraphs
    recommendations = summary.get("recommendations", [])

    rec_slots = [(93, 94), (95, 96), (97, 98), (99, 100)]

    for idx, rec in enumerate(recommendations):
        if idx >= len(rec_slots):
            break
        heading_idx, body_idx = rec_slots[idx]
        title = rec.get("title", "")
        body = rec.get("body", "")

        if heading_idx < len(paragraphs) and title:
            _replace_paragraph_text(paragraphs[heading_idx], title, font_size_pt=16, bold=True)
        if body_idx < len(paragraphs) and body:
            _replace_paragraph_text_preserve_format(paragraphs[body_idx], body)


# ── Chapter 5: Attachment ────────────────────────────────────────────────────

def _replace_chapter5(doc: Document, attachment: dict):
    """Replace 附件-问卷题目内容 section."""
    questions = attachment.get("questions", [])
    paragraphs = doc.paragraphs
    heading = next((p for p in paragraphs if p.text.strip() in {
        "附件-问卷题目内容", "五、附件-问卷题目内容"
    }), None)
    if heading is None:
        raise ValueError("Template missing Chapter 5 attachment heading")

    question_style = copy.deepcopy(paragraphs[113]._p.pPr)
    option_style = copy.deepcopy(paragraphs[114]._p.pPr)
    node = heading._p.getnext()
    while node is not None and node.tag != f"{{{W_NS}}}sectPr":
        next_node = node.getnext()
        node.getparent().remove(node)
        node = next_node

    for q_idx, question in enumerate(questions, start=1):
        q_para = doc.add_paragraph()
        if q_para._p.pPr is not None:
            q_para._p.remove(q_para._p.pPr)
        q_para._p.insert(0, copy.deepcopy(question_style))
        q_run = q_para.add_run(f"{question.get('number', q_idx)}.{question.get('question', '')}")
        q_run.font.name = "宋体"
        q_run.font.size = Pt(16)
        q_run.bold = True
        for opt_idx, option in enumerate(question.get("options", [])):
            code = str(option.get("code") or option.get("label") or chr(65 + opt_idx)).strip()
            text = str(option.get("text", "")).strip()
            text = text.removeprefix(f"{code}.").removeprefix(f"{code}、").strip()
            opt_para = doc.add_paragraph()
            if opt_para._p.pPr is not None:
                opt_para._p.remove(opt_para._p.pPr)
            opt_para._p.insert(0, copy.deepcopy(option_style))
            opt_run = opt_para.add_run(f"{code}.{text}")
            opt_run.font.name = "宋体"
            opt_run.font.size = Pt(14)
            opt_run.bold = False


def _format_primary_headings(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text not in PRIMARY_HEADINGS:
            continue
        _replace_paragraph_text(paragraph, PRIMARY_HEADINGS[text], font_size_pt=20, bold=True)
        p_pr = paragraph._p.get_or_add_pPr()
        outline = p_pr.find(f"{{{W_NS}}}outlineLvl")
        if outline is None:
            outline = etree.SubElement(p_pr, f"{{{W_NS}}}outlineLvl")
        outline.set(f"{{{W_NS}}}val", "0")


# ── TOC insertion ────────────────────────────────────────────────────────────

def _insert_toc_field(docx_path: Path, insert_after_para_idx: int = TOC_INSERT_AFTER):
    """Insert a TOC field into the DOCX using direct XML manipulation.

    This must be called AFTER python-docx saves the file, since python-docx
    doesn't support field codes natively.
    """
    # Read the DOCX as ZIP
    with zipfile.ZipFile(docx_path, "r") as zf:
        file_contents = {}
        for item in zf.namelist():
            file_contents[item] = zf.read(item)

    # Parse document.xml
    doc_xml = etree.fromstring(file_contents["word/document.xml"])
    body = doc_xml.find(f"{{{W_NS}}}body")
    paragraphs = body.findall(f"{{{W_NS}}}p")

    if insert_after_para_idx >= len(paragraphs):
        print(f"WARNING: Cannot insert TOC at paragraph {insert_after_para_idx}, only {len(paragraphs)} paragraphs")
        return

    # Create TOC paragraph with field code
    toc_para = etree.SubElement(body, f"{{{W_NS}}}p")
    # Move it after the target paragraph
    target_para = paragraphs[insert_after_para_idx]
    target_para.addnext(toc_para)

    # Add paragraph properties (centered, with spacing)
    pPr = etree.SubElement(toc_para, f"{{{W_NS}}}pPr")
    jc = etree.SubElement(pPr, f"{{{W_NS}}}jc")
    jc.set(f"{{{W_NS}}}val", "center")

    # Add TOC field: begin
    r1 = etree.SubElement(toc_para, f"{{{W_NS}}}r")
    fldChar1 = etree.SubElement(r1, f"{{{W_NS}}}fldChar")
    fldChar1.set(f"{{{W_NS}}}fldCharType", "begin")

    # Add TOC field: instruction
    r2 = etree.SubElement(toc_para, f"{{{W_NS}}}r")
    instrText = etree.SubElement(r2, f"{{{W_NS}}}instrText")
    instrText.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instrText.text = ' TOC \\o "1-2" \\h \\z \\u '

    # Add TOC field: separate
    r3 = etree.SubElement(toc_para, f"{{{W_NS}}}r")
    fldChar2 = etree.SubElement(r3, f"{{{W_NS}}}fldChar")
    fldChar2.set(f"{{{W_NS}}}fldCharType", "separate")

    # Add TOC field: placeholder text
    r4 = etree.SubElement(toc_para, f"{{{W_NS}}}r")
    t = etree.SubElement(r4, f"{{{W_NS}}}t")
    t.text = "目录（请在 Word 中按 Ctrl+A, F9 更新）"

    # Add TOC field: end
    r5 = etree.SubElement(toc_para, f"{{{W_NS}}}r")
    fldChar3 = etree.SubElement(r5, f"{{{W_NS}}}fldChar")
    fldChar3.set(f"{{{W_NS}}}fldCharType", "end")

    # Also add a page break after TOC
    pb_para = etree.SubElement(body, f"{{{W_NS}}}p")
    target_para.addnext(pb_para)  # This goes after target, before TOC
    # Actually, let's put it after TOC
    toc_para.addnext(pb_para)
    pPr_pb = etree.SubElement(pb_para, f"{{{W_NS}}}pPr")
    # We need to re-add toc_para after pb_para
    # Actually, let me redo this more carefully

    # Force every textual run, including cover text boxes and the TOC
    # placeholder, to use explicit Songti declarations. Keep the template's
    # existing size, bold, color, spacing, and paragraph properties unchanged.
    for run in doc_xml.findall(f".//{{{W_NS}}}r"):
        if run.find(f".//{{{W_NS}}}t") is None:
            continue
        r_pr = run.find(f"{{{W_NS}}}rPr")
        if r_pr is None:
            r_pr = etree.Element(f"{{{W_NS}}}rPr")
            run.insert(0, r_pr)
        r_fonts = r_pr.find(f"{{{W_NS}}}rFonts")
        if r_fonts is None:
            r_fonts = etree.Element(f"{{{W_NS}}}rFonts")
            r_pr.insert(0, r_fonts)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            r_fonts.set(f"{{{W_NS}}}{attr}", "宋体")
        for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
            r_fonts.attrib.pop(f"{{{W_NS}}}{attr}", None)

    # Write back
    file_contents["word/document.xml"] = etree.tostring(
        doc_xml, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    # Word regenerates TOC entries from styles, so explicit run formatting in
    # document.xml alone is insufficient. Set every document style's font to
    # Songti while retaining its template-defined size and emphasis.
    if "word/styles.xml" in file_contents:
        styles_xml = etree.fromstring(file_contents["word/styles.xml"])
        style_owners = list(styles_xml.findall(f".//{{{W_NS}}}style"))
        default_r_pr = styles_xml.find(f".//{{{W_NS}}}docDefaults/{{{W_NS}}}rPrDefault/{{{W_NS}}}rPr")
        for owner in style_owners:
            r_pr = owner.find(f"{{{W_NS}}}rPr")
            if r_pr is None:
                r_pr = etree.SubElement(owner, f"{{{W_NS}}}rPr")
            _set_songti_rfonts(r_pr)
        if default_r_pr is not None:
            _set_songti_rfonts(default_r_pr)
        file_contents["word/styles.xml"] = etree.tostring(
            styles_xml, xml_declaration=True, encoding="UTF-8", standalone=True
        )

    # Update settings.xml to set updateFields=true
    if "word/settings.xml" in file_contents:
        settings_xml = etree.fromstring(file_contents["word/settings.xml"])
        # Check if updateFields already exists
        existing = settings_xml.find(f"{{{W_NS}}}updateFields")
        if existing is None:
            # Add it as first child
            uf = etree.SubElement(settings_xml, f"{{{W_NS}}}updateFields")
            uf.set(f"{{{W_NS}}}val", "true")
            # Move to first position
            settings_xml.insert(0, uf)
        else:
            existing.set(f"{{{W_NS}}}val", "true")
        file_contents["word/settings.xml"] = etree.tostring(
            settings_xml, xml_declaration=True, encoding="UTF-8", standalone=True
        )

    # Write the modified DOCX
    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, data in file_contents.items():
            zf.writestr(name, data)


def _set_songti_rfonts(r_pr: etree._Element) -> None:
    r_fonts = r_pr.find(f"{{{W_NS}}}rFonts")
    if r_fonts is None:
        r_fonts = etree.Element(f"{{{W_NS}}}rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(f"{{{W_NS}}}{attr}", "宋体")
    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        r_fonts.attrib.pop(f"{{{W_NS}}}{attr}", None)


# ── Main render function ─────────────────────────────────────────────────────

def render_from_template(
    payload: dict,
    template_path: Path,
    output_or_manifest_path: Path,
    legacy_output_path: Optional[Path] = None,
) -> Path:
    """Render a patient report from payload v2 using the customer DOCX template.

    Args:
        payload: Payload v2 dict with meta, intro, chapter2, feedback, summary, attachment.
        template_path: Path to the customer DOCX template.
        output_or_manifest_path: Output DOCX path, or the legacy manifest path.
        legacy_output_path: Legacy fourth positional argument for output DOCX.

    Returns:
        Path to the output DOCX.
    """
    # Preserve the previous public call shape:
    # render_from_template(payload, template, manifest, output).
    output_path = Path(legacy_output_path or output_or_manifest_path)

    # Validate payload structure
    required_keys = ["meta", "intro", "chapter2", "feedback", "summary", "attachment"]
    missing = [k for k in required_keys if k not in payload]
    if missing:
        raise ValueError(f"Payload v2 missing required keys: {', '.join(missing)}")

    items = payload["chapter2"].get("items", [])
    if len(items) > 11:
        raise ValueError(f"Payload has {len(items)} questions, maximum is 11")

    # Step 1: Copy template
    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template_path, output_path)

    # Step 2: Load with python-docx for text replacement
    doc = Document(str(output_path))

    # Step 3: Replace cover
    _replace_cover(doc, payload.get("meta", {}))

    # Step 4: Replace chapter 1 (introduction)
    _replace_chapter1(doc, payload.get("intro", {}))

    # Step 5: Replace chapter 2 (data analysis)
    _replace_chapter2(doc, payload.get("chapter2", {}))

    # Step 6: Replace chapter 3 (feedback)
    _replace_chapter3(doc, payload.get("feedback", {}))

    # Step 7: Replace chapter 4 (summary)
    _replace_chapter4(doc, payload.get("summary", {}))

    # Step 8: Replace chapter 5 (attachment)
    _replace_chapter5(doc, payload.get("attachment", {}))

    # Step 9: Apply numbered primary headings and outline levels.
    _format_primary_headings(doc)

    # Step 10: Save (python-docx)
    doc.save(str(output_path))

    # Step 11: Insert TOC + updateFields (direct XML manipulation)
    _insert_toc_field(output_path, TOC_INSERT_AFTER)

    # Step 12: Update native charts
    chart_slots = payload_to_chart_slots(payload)
    if chart_slots:
        _, chart_results = update_charts_in_docx(output_path, chart_slots, output_path)
        failures = [result for result in chart_results if not result.ok]
        for r in chart_results:
            if not r.ok:
                print(f"  WARNING: Chart slot {r.slot_number} has issues:")
                for e in r.errors:
                    print(f"    {e}")
        if failures:
            raise ValueError(f"Native chart update failed for {len(failures)} slot(s)")

    # Step 13: Validate output
    _validate_output(output_path)

    return output_path


def _validate_output(docx_path: Path):
    """Run basic validation on the output DOCX."""
    try:
        doc = Document(str(docx_path))
        # Basic checks
        assert len(doc.paragraphs) > 100, f"Too few paragraphs: {len(doc.paragraphs)}"

        # Check for TOC field (via XML)
        with zipfile.ZipFile(docx_path) as zf:
            doc_xml = etree.fromstring(zf.read("word/document.xml"))
            # Look for fldChar with fldCharType="begin" followed by instrText containing "TOC"
            fld_chars = doc_xml.findall(f".//{{{W_NS}}}fldChar")
            has_toc = False
            for fc in fld_chars:
                if fc.get(f"{{{W_NS}}}fldCharType") == "begin":
                    # Check if next sibling r contains instrText with TOC
                    parent_r = fc.getparent()
                    next_r = parent_r.getnext()
                    if next_r is not None:
                        instr = next_r.find(f"{{{W_NS}}}instrText")
                        if instr is not None and "TOC" in (instr.text or ""):
                            has_toc = True
                            break
            assert has_toc, "No TOC field found in output"

            # Check updateFields
            if "word/settings.xml" in zf.namelist():
                settings_xml = etree.fromstring(zf.read("word/settings.xml"))
                uf = settings_xml.find(f"{{{W_NS}}}updateFields")
                assert uf is not None, "updateFields not found in settings.xml"
                assert uf.get(f"{{{W_NS}}}val") == "true", f"updateFields is not 'true'"

            # Check for charts
            chart_count = len([n for n in zf.namelist() if n.startswith("word/charts/chart") and n.endswith(".xml")])
            assert chart_count == 11, f"Expected 11 charts, found {chart_count}"

            # Check for workbooks
            wb_count = len([n for n in zf.namelist() if "Workbook" in n and n.endswith(".xlsx")])
            assert wb_count == 11, f"Expected 11 workbooks, found {wb_count}"

        print(f"✓ Output validation passed: {docx_path}")

    except Exception as e:
        raise ValueError(f"Output validation failed: {e}") from e


# ── CLI ──────────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Render patient report from payload v2 using customer DOCX template."
    )
    parser.add_argument("payload_json", help="Path to payload v2 JSON file")
    parser.add_argument("-o", "--output", required=True, help="Output DOCX path")
    parser.add_argument("--template", help="Template DOCX path (default: templates/patient-unified-v1.docx)")
    parser.add_argument("--manifest", help="Manifest JSON path (optional)")
    args = parser.parse_args()

    # Load payload
    payload = json.loads(Path(args.payload_json).read_text(encoding="utf-8"))

    # Resolve template path
    if args.template:
        template_path = Path(args.template)
    else:
        template_path = PROJECT_ROOT / "templates" / "patient-unified-v1.docx"

    if not template_path.exists():
        print(f"ERROR: Template not found: {template_path}")
        sys.exit(1)

    output_path = Path(args.output)
    manifest_path = Path(args.manifest) if args.manifest else None

    # Render
    if manifest_path is not None:
        result_path = render_from_template(payload, template_path, manifest_path, output_path)
    else:
        result_path = render_from_template(payload, template_path, output_path)
    print(result_path)


if __name__ == "__main__":
    main()
